import os
import aiofiles
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import aiohttp
from dotenv import load_dotenv
import re
import asyncio
from copy import deepcopy

load_dotenv()

class TranslationService:
    def __init__(self):
        self.openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
        if not self.openrouter_api_key:
            raise ValueError("OPENROUTER_API_KEY not found in environment variables")
        
        self.max_tokens_per_chunk = 1000  # Increased for faster processing
        self.max_retries = 3
        self.retry_delay = 1
        self.batch_size = 5  # Number of paragraphs to process in parallel

    async def send_progress(self, client_id: str, stage: str, progress: int):
        """Send progress updates via WebSocket"""
        if client_id:
            from app.main import manager
            try:
                message = {
                    "type": "progress",
                    "stage": stage,
                    "progress": progress,
                    "message": stage
                }
                print(f"Sending progress update: {message}")  # Debug log
                await manager.send_progress(client_id, message)
                print(f"Progress: {stage} - {progress}%")  # Terminal progress
            except Exception as e:
                print(f"Error sending progress update: {str(e)}")
                # Try to send error to client
                try:
                    await manager.send_progress(client_id, {
                        "type": "error",
                        "message": f"Error during translation: {str(e)}"
                    })
                except:
                    pass

    def copy_run_format(self, source_run, target_run):
        """Copy all formatting from source run to target run"""
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        target_run.font.color.rgb = source_run.font.color.rgb if source_run.font.color.rgb else RGBColor(0, 0, 0)
        target_run.style = source_run.style

    def copy_paragraph_format(self, source_paragraph, target_paragraph):
        """Copy all formatting from source paragraph to target paragraph"""
        target_paragraph.alignment = source_paragraph.alignment
        target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
        target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
        target_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
        target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
        target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
        target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
        target_paragraph.style = source_paragraph.style

    def ensure_proper_spacing(self, text: str) -> str:
        """Ensure proper spacing in translated text while preserving original spacing patterns"""
        # Remove extra spaces between words
        text = re.sub(r'\s{2,}', ' ', text)

        # Ensure proper spacing around punctuation
        text = re.sub(r'\s*([.,!?;:])\s*', r'\1 ', text)

        # Remove spaces before punctuation
        text = re.sub(r'\s+([.,!?;:])', r'\1', text)

        # Ensure a single space after punctuation
        text = re.sub(r'([.,!?;:])([^\s])', r'\1 \2', text)

        # Trim leading and trailing spaces
        text = text.strip()

        return text
    
    def split_text_into_chunks(self, text: str) -> list:
        """Split text into smaller chunks while preserving sentence structure"""
        if not text:
            return []
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_tokens = len(sentence) // 4
            if current_length + sentence_tokens > self.max_tokens_per_chunk and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_length = sentence_tokens
            else:
                current_chunk.append(sentence)
                current_length += sentence_tokens
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    async def translate_chunk_with_retry(self, session, chunk: str, target_language: str) -> str:
        """Translate a single chunk with retry logic"""
        # Handle empty chunks
        if not chunk or not chunk.strip():
            return chunk
            
        # Map target language (e.g., "Modern English" → "Contemporary English")
        actual_target_language = self.get_actual_target_language(target_language)
            
        # Preserve leading and trailing whitespace
        leading_space = ""
        trailing_space = ""
        
        if chunk.startswith(" "):
            leading_space = " " * (len(chunk) - len(chunk.lstrip()))
            chunk = chunk.lstrip()
            
        if chunk.endswith(" "):
            trailing_space = " " * (len(chunk) - len(chunk.rstrip()))
            chunk = chunk.rstrip()
            
        for attempt in range(self.max_retries):
            try:
                headers = {
                    "Authorization": f"Bearer {self.openrouter_api_key}",
                    "Content-Type": "application/json",
                    "HTTP-Referer": "https://book-translator.com.equationofintelligence.com"
                }
                
                # Get the appropriate system prompt for this translation
                system_prompt = self.get_system_prompt(actual_target_language)
                
                payload = {
                    "model": "openai/gpt-4o",  # Using GPT-4 for better translation
                    "messages": [
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": f"Translate this text to {actual_target_language}: {chunk}"
                        }
                    ],
                    "max_tokens": self.max_tokens_per_chunk * 2,  # Increased token limit for translations
                    "temperature": 0.3,  # Balanced between creativity and accuracy
                    "stream": False,
                    "top_p": 0.9  # Added for better translation quality
                }

                print(f"Translating text to {actual_target_language}: {chunk}")  # Debug log
                async with session.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=30
                ) as response:
                    if response.status == 200:
                        result = await response.json()
                        translated_text = result["choices"][0]["message"]["content"]
                        
                        print(f"Translation result: {translated_text}")  # Debug log
                        
                        # Verify that the text was actually translated
                        if translated_text.strip() == chunk.strip():
                            print(f"Warning: Translation result is identical to input")
                            
                        # Restore leading and trailing whitespace
                        translated_text = leading_space + translated_text + trailing_space
                        
                        # Make sure there's proper spacing between words
                        translated_text = self.ensure_proper_spacing(translated_text)
                        
                        return translated_text
                    elif response.status == 402:
                        error_data = await response.json()
                        raise Exception(f"Token limit exceeded: {error_data.get('error', {}).get('message', 'Unknown error')}")
                    else:
                        error_text = await response.text()
                        raise Exception(f"API request failed with status {response.status}: {error_text}")
                        
            except Exception as e:
                if attempt == self.max_retries - 1:
                    raise
                await asyncio.sleep(self.retry_delay * (attempt + 1))

    def copy_document_settings(self, source_doc, target_doc):
        """
        This method is no longer used in the new translation approach,
        since we make an exact file copy instead of trying to recreate the document.
        Kept for compatibility with other code that might call this.
        """
        # Copy section settings for the first section
        try:
            target_doc.sections[0].page_height = source_doc.sections[0].page_height
            target_doc.sections[0].page_width = source_doc.sections[0].page_width
            target_doc.sections[0].left_margin = source_doc.sections[0].left_margin
            target_doc.sections[0].right_margin = source_doc.sections[0].right_margin
            target_doc.sections[0].top_margin = source_doc.sections[0].top_margin
            target_doc.sections[0].bottom_margin = source_doc.sections[0].bottom_margin
            target_doc.sections[0].header_distance = source_doc.sections[0].header_distance
            target_doc.sections[0].footer_distance = source_doc.sections[0].footer_distance
        except Exception as e:
            print(f"Warning: Could not copy section properties: {str(e)}")
            
        # Copy styles
        try:
            for style in source_doc.styles:
                if style.name not in target_doc.styles:
                    try:
                        target_doc.styles.add_style(
                            style.name,
                            style.type,
                            style.base_style
                        )
                    except Exception as style_err:
                        print(f"Warning: Could not copy style {style.name}: {str(style_err)}")
        except Exception as e:
            print(f"Warning: Error copying styles: {str(e)}")

    async def translate_paragraph_word_by_word(self, paragraph, session, target_language: str) -> str:
        """Translate a paragraph word by word while preserving formatting"""
        translated_text = []
        for run in paragraph.runs:
            words = run.text.split()  # Split text into words
            translated_words = []
            for word in words:
                try:
                    # Translate each word
                    translated_word = await self.translate_chunk_with_retry(session, word, target_language)
                    translated_words.append(translated_word)
                except Exception as e:
                    print(f"Error translating word '{word}': {str(e)}")
                    translated_words.append(word)  # Keep original word if translation fails
            # Reconstruct the run text
            translated_text.append(" ".join(translated_words))
        return " ".join(translated_text)

    async def process_document(self, input_path: str, target_language: str, client_id: str = None) -> str:
        """Process a document file for translation while preserving formatting"""
        try:
            # Initial progress update
            await self.send_progress(client_id, "Starting translation process...", 0)

            # Create output filename
            base_name, ext = os.path.splitext(input_path)
            
            # Map the target language for the filename (e.g., modern_english → contemporary_english)
            filename_language = target_language.lower()
            if filename_language in ["modern english", "modern_english"]:
                filename_language = "contemporary_english"
            else:
                filename_language = filename_language.replace(" ", "_")
                
            output_path = f"{base_name}_{filename_language}_translated{ext}"
            
            # First, make a direct file copy to preserve all document properties
            import shutil
            shutil.copy2(input_path, output_path)
            
            # Now open the copied document for modification
            await self.send_progress(client_id, "Loading document...", 5)
            doc = Document(output_path)
            
            # Get all content-containing elements that need translation
            await self.send_progress(client_id, "Analyzing document structure...", 10)
            
            # Create lists of all elements that may contain text
            paragraphs = []
            
            # Store tables separately for specialized processing
            tables = list(doc.tables)
            
            # Main document paragraphs
            paragraphs.extend(doc.paragraphs)
            
            # Table cell paragraphs - we'll process these with table handling
            # for table in doc.tables:
            #     for row in table.rows:
            #         for cell in row.cells:
            #             paragraphs.extend(cell.paragraphs)
            
            # Headers and footers
            for section in doc.sections:
                # Header paragraphs
                paragraphs.extend(section.header.paragraphs)
                # Footer paragraphs
                paragraphs.extend(section.footer.paragraphs)
            
            # Count total paragraphs that actually have text content
            total_paragraphs = len([p for p in paragraphs if p.text.strip()])
            translated_count = 0
            
            # Process all main paragraphs (not including tables)
            async with aiohttp.ClientSession() as session:
                # First translate regular paragraphs
                for i, paragraph in enumerate(paragraphs):
                    if not paragraph.text.strip():
                        continue  # Skip empty paragraphs
                        
                    # Update progress (20-70%)
                    translated_count += 1
                    progress = 20 + int((translated_count / total_paragraphs) * 60)
                    
                    await self.send_progress(
                        client_id,
                        f"Translating paragraph {translated_count} of {total_paragraphs}...",
                        progress
                    )
                    
                    # Collect runs with text for the entire paragraph to preserve spacing context
                    paragraph_text = paragraph.text
                    
                    # Skip if paragraph is empty or contains special fields
                    if not paragraph_text.strip() or ('{' in paragraph_text and '}' in paragraph_text):
                        continue
                        
                    # First check if this paragraph needs capitalization fixing
                    needs_capitalization_fix = paragraph_text.isupper() and len(paragraph_text.split()) > 1
                        
                    try:
                        # Use the new paragraph processing method
                        await self.process_paragraph(paragraph, session, target_language)
                            
                    except Exception as e:
                        print(f"Error translating paragraph: {str(e)}")
                        # Keep original text if translation fails
                        continue
            
            # Now process tables with specialized handling
            if tables:
                await self.send_progress(client_id, "Processing tables...", 80)
                
                # Process each table to preserve structure
                for table_idx, table in enumerate(tables):
                    progress = 80 + int((table_idx / len(tables)) * 15)
                    
                    await self.send_progress(
                        client_id,
                        f"Translating table {table_idx+1} of {len(tables)}...",
                        progress
                    )
                    
                    # Use our specialized table handling method
                    await self.handle_table_translation(session, table, target_language, client_id)
            
            # Save the document (already saved to the output path)
            await self.send_progress(client_id, "Finalizing document...", 95)
            
            try:
                # Save with proper error handling
                doc.save(output_path)
                
                # Verify the document can be opened
                test_doc = Document(output_path)
                
                # Send completion update
                await self.send_progress(client_id, "Translation complete!", 100)
                return output_path
                
            except Exception as save_error:
                print(f"Error saving document: {str(save_error)}")
                await self.send_progress(client_id, f"Error saving: {str(save_error)}", -1)
                raise
                
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            await self.send_progress(client_id, f"Error: {str(e)}", -1)
            raise

    async def translate_paragraph(self, paragraph, session, target_language: str) -> str:
        """Translate a paragraph while preserving formatting"""
        # Get text from runs
        runs_text = []
        for run in paragraph.runs:
            if run.text.strip():
                runs_text.append(run.text)

        if not runs_text:
            return []

        # Translate each text segment
        translated_texts = []
        for text in runs_text:
            if text.strip():
                try:
                    translated_text = await self.translate_chunk_with_retry(session, text, target_language)
                    translated_texts.append(translated_text)
                except Exception as e:
                    print(f"Error translating text: {str(e)}")
                    translated_texts.append(text)  # Keep original text if translation fails
            else:
                translated_texts.append(text)

        return translated_texts

    async def copy_paragraph_with_translation(self, source_paragraph, target_doc, session, target_language: str):
        """
        This method is kept for compatibility but is no longer used in the new translation approach.
        The new approach modifies runs in-place instead of copying paragraphs.
        """
        try:
            # Create a new paragraph
            new_paragraph = target_doc.add_paragraph()
            
            # Copy paragraph formatting
            self.copy_paragraph_format(source_paragraph, new_paragraph)
            
            # Translate and preserve runs
            translated_texts = await self.translate_paragraph(source_paragraph, session, target_language)
            
            # Add translated runs with their original formatting
            for i, (run, translated_text) in enumerate(zip(source_paragraph.runs, translated_texts)):
                if run.text.strip():  # Only process non-empty runs
                    new_run = new_paragraph.add_run(translated_text)
                    self.copy_run_format(run, new_run)
                else:
                    # Preserve empty runs for spacing/formatting
                    new_run = new_paragraph.add_run(run.text)
                    self.copy_run_format(run, new_run)
                    
            return new_paragraph
            
        except Exception as e:
            print(f"Error in copy_paragraph_with_translation: {str(e)}")
            # Create a basic paragraph with the original text as fallback
            return target_doc.add_paragraph(source_paragraph.text)

    def preserve_capitalization(self, original_text: str, translated_text: str) -> str:
        """Preserve the capitalization pattern from the original text in the translated text"""
        if not original_text or not translated_text:
            return translated_text
            
        # If original text is all uppercase, check if it should be converted to title case
        if original_text.isupper() and len(original_text.split()) > 1:
            # Check if this looks like a title rather than an acronym or short emphasis
            words = original_text.split()
            if len(words) > 1 and any(len(word) > 3 for word in words):
                # This is likely a title or heading in all caps, convert to title case
                return self.to_title_case(translated_text)
            else:
                # Keep all uppercase for acronyms, short emphases, etc.
                return translated_text.upper()
            
        # If original text is all uppercase but looks like an acronym (short word), keep it uppercase
        elif original_text.isupper():
            return translated_text.upper()
            
        # If original text is all lowercase, make translated text all lowercase
        elif original_text.islower():
            return translated_text.lower()
            
        # If first character is uppercase, capitalize the first character of translated text
        elif original_text[0].isupper() and translated_text:
            return translated_text[0].upper() + translated_text[1:]
                
        return translated_text
        
    def to_title_case(self, text: str) -> str:
        """Convert text to title case, keeping small words lowercase except at the beginning"""
        if not text:
            return text
            
        # Split the text into words
        words = text.lower().split()
        
        # List of small words that shouldn't be capitalized in titles
        # unless they're the first or last word
        small_words = {'a', 'an', 'the', 'and', 'but', 'or', 'for', 'nor', 'on', 'at', 
                      'to', 'from', 'by', 'in', 'of', 'with', 'as'}
        
        # Capitalize the first word and any word that's not in the small_words list
        result = []
        for i, word in enumerate(words):
            if i == 0 or i == len(words) - 1 or word not in small_words:
                result.append(word.capitalize())
            else:
                result.append(word)
                
        return ' '.join(result)
        
    def distribute_translated_text(self, runs, translated_text: str) -> None:
        """Distribute translated text across multiple runs while preserving formatting and colors"""
        # Get only runs that contain text
        original_runs = [run for run in runs if run.text.strip()]
        if not original_runs:
            return
            
        original_text_total = sum(len(run.text) for run in original_runs)
        if original_text_total == 0:
            return
            
        # Store the original run properties for later reference
        run_info = []
        for run in original_runs:
            original_text = run.text
            run_info.append({
                'run': run,
                'original_text': original_text,
                'length': len(original_text),
                'leading_spaces': len(original_text) - len(original_text.lstrip()),
                'trailing_spaces': len(original_text) - len(original_text.rstrip()),
                'is_capitalized': original_text.strip() and original_text.strip()[0].isupper() if original_text.strip() else False,
                'is_uppercase': original_text.isupper() if original_text.strip() else False,
                'is_lowercase': original_text.islower() if original_text.strip() else False,
                'proportion': len(original_text) / original_text_total
            })
        
        # Distribute the translated text across runs
        remaining_text = translated_text
        
        for i, info in enumerate(run_info):
            if not remaining_text:
                info['run'].text = ""
                continue
                
            # For the last run, use all remaining text
            if i == len(run_info) - 1:
                new_text = remaining_text
            else:
                # Calculate proportional length for this run
                target_length = max(1, int(len(translated_text) * info['proportion']))
                
                # Try to break at word boundaries when possible
                if len(remaining_text) > target_length:
                    space_pos = remaining_text.rfind(' ', 0, target_length)
                    if space_pos == -1 or space_pos < target_length / 2:
                        # If no good word boundary, use the calculated length
                        new_text = remaining_text[:target_length]
                    else:
                        # Break at word boundary
                        new_text = remaining_text[:space_pos+1]
                else:
                    new_text = remaining_text
            
            # Preserve original capitalization pattern
            if info['is_uppercase'] and len(info['original_text'].split()) > 1:
                # Check if this is a title that should be converted to title case
                words = info['original_text'].split()
                if any(len(word) > 3 for word in words):
                    new_text = self.to_title_case(new_text)
                else:
                    new_text = new_text.upper()
            elif info['is_uppercase']:
                new_text = new_text.upper()
            elif info['is_lowercase']:
                new_text = new_text.lower()
            elif info['is_capitalized'] and new_text:
                new_text = new_text[0].upper() + new_text[1:]
            
            # Preserve original spacing patterns
            if info['leading_spaces'] > 0 or info['trailing_spaces'] > 0:
                content = new_text.strip()
                new_text = (' ' * info['leading_spaces']) + content + (' ' * info['trailing_spaces'])
            
            # Update the run text and remove the used portion from remaining text
            info['run'].text = new_text
            remaining_text = remaining_text[len(new_text):]
        
        # If we have any remaining text, add it to the last run
        if remaining_text and run_info:
            run_info[-1]['run'].text += remaining_text

    async def handle_table_translation(self, session, table, target_language: str, client_id: str = None) -> None:
        """Translate text within a table while preserving structure, spacing, colors and formatting"""
        try:
            # Process each cell in the table
            for row in table.rows:
                for cell in row.cells:
                    # Process each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        # Skip empty paragraphs
                        if not paragraph.text.strip():
                            continue
                            
                        try:
                            # Check if paragraph needs capitalization fixing
                            needs_capitalization_fix = paragraph.text.isupper() and len(paragraph.text.split()) > 1
                            
                            # Translate the paragraph text
                            translated_text = await self.translate_chunk_with_retry(
                                session,
                                paragraph.text,
                                target_language
                            )
                            
                            # Apply the same formatting preservation logic as regular paragraphs
                            if len(paragraph.runs) == 1:
                                # Simple case with one run
                                original_text = paragraph.runs[0].text
                                
                                # Fix capitalization issues
                                if original_text.isupper() and len(original_text.split()) > 1:
                                    # Check if this is a title that should be converted to title case
                                    words = original_text.split()
                                    if any(len(word) > 3 for word in words):
                                        translated_text = self.to_title_case(translated_text)
                                    else:
                                        translated_text = self.preserve_capitalization(original_text, translated_text)
                                else:
                                    translated_text = self.preserve_capitalization(original_text, translated_text)
                                
                                # Preserve original spacing
                                leading_spaces = len(original_text) - len(original_text.lstrip())
                                trailing_spaces = len(original_text) - len(original_text.rstrip())
                                
                                if leading_spaces > 0 or trailing_spaces > 0:
                                    content = translated_text.strip()
                                    translated_text = (' ' * leading_spaces) + content + (' ' * trailing_spaces)
                                
                                paragraph.runs[0].text = translated_text
                            else:
                                # Complex case with multiple runs and formatting
                                self.distribute_translated_text(paragraph.runs, translated_text)
                                
                        except Exception as cell_error:
                            print(f"Error translating table cell: {str(cell_error)}")
                            # Continue with next paragraph, keeping original text
                            continue
        except Exception as table_error:
            print(f"Error processing table: {str(table_error)}")
            # Continue processing the document even if a table fails

    def fix_all_caps_text(self, text: str) -> str:
        """Fix text that is in all caps by converting it to more natural sentence case or title case
        
        This method examines text and converts problematic all caps sections to more
        readable formats while preserving legitimate uses of uppercase (acronyms, etc.)
        """
        if not text or not any(c.isupper() for c in text):
            return text  # No uppercase characters, nothing to do
            
        # If the entire text is uppercase and it's clearly a sentence or paragraph
        if text.isupper() and len(text.split()) > 3:
            # Convert to sentence case (first letter capitalized, rest lowercase)
            sentences = re.split(r'([.!?]\s+)', text.lower())
            result = ""
            for i in range(0, len(sentences), 2):
                sentence = sentences[i]
                # Capitalize first letter of each sentence
                if sentence and sentence[0].isalpha():
                    sentence = sentence[0].upper() + sentence[1:]
                result += sentence
                # Add the separator back
                if i+1 < len(sentences):
                    result += sentences[i+1]
            return result
            
        # If it's a short title (likely a heading), convert to title case
        elif text.isupper() and 1 < len(text.split()) <= 10:
            return self.to_title_case(text)
            
        # Special case: mixed case with ALL CAPS sections
        # This is more complex and would require identifying ALL CAPS phrases within text
        # For simplicity, we're only handling full sentences/paragraphs
            
        return text

    def get_actual_target_language(self, target_language: str) -> str:
        """Map requested target language to actual translation language
        
        This is used to handle special cases like Modern English → Contemporary English
        """
        # Map "modern english" to "contemporary english"
        if target_language.lower() in ["modern english", "modern_english"]:
            return "Contemporary English"
            
        # Return the original target language for all other cases
        return target_language

    def get_system_prompt(self, target_language: str) -> str:
        """Get the appropriate system prompt based on the target language"""
        # Special case for Contemporary English
        if target_language.lower() == "contemporary english":
            return (
                f"You are a professional translator specializing in contemporary language. "
                f"Translate the following text to Contemporary English - meaning clear, natural, "
                f"modern-day English as used by educated native speakers today. "
                f"Use current expressions, vocabulary, and sentence structures. "
                f"Avoid archaic terms, outdated expressions, and complicated syntax. "
                f"IMPORTANT: Preserve the exact spacing and formatting. "
                f"For capitalization: 1) If the text is ALL CAPS and appears to be a normal sentence "
                f"or paragraph, convert it to normal sentence case; 2) If the text is ALL CAPS and "
                f"appears to be a title or heading, convert it to Title Case; 3) For acronyms or "
                f"intentionally emphasized short phrases in ALL CAPS, keep them in ALL CAPS. "
                f"Preserve all special characters and punctuation."
            )
            
        # Default prompt for other languages
        return (
            f"You are a professional translator. Your task is to translate the following text into {target_language}. "
            f"Do not keep the original English text - translate everything into {target_language}. "
            f"IMPORTANT RULES:\n"
            f"1. Translate ALL text into {target_language}\n"
            f"2. Keep the same formatting and spacing\n"
            f"3. For capitalization:\n"
            f"   - Normal sentences: use normal case in {target_language}\n"
            f"   - Titles and headings: use title case in {target_language}\n"
            f"   - Acronyms: keep in ALL CAPS\n"
            f"4. Preserve all special characters and punctuation\n"
            f"5. The entire response should be in {target_language} only"
        )
    
    def preserve_run_properties(self, run):
        """Store all properties of a run for preservation"""
        properties = {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font': {
                'name': run.font.name,
                'size': run.font.size,
                'color': run.font.color.rgb if run.font.color else None
            },
            'style': run.style
        }
        
        # Try to get XML level properties only if they exist
        try:
            if hasattr(run._element, 'rPr') and run._element.rPr is not None:
                rPr = run._element.rPr
                if hasattr(rPr, 'spacing'):
                    properties['spacing'] = rPr.spacing
                if hasattr(rPr, 'position'):
                    properties['position'] = rPr.position
        except:
            pass
            
        return properties

    def restore_run_properties(self, run, properties):
        """Restore all properties to a run"""
        run.bold = properties['bold']
        run.italic = properties['italic']
        run.underline = properties['underline']
        run.font.name = properties['font']['name']
        if properties['font']['size']:
            run.font.size = properties['font']['size']
        if properties['font']['color']:
            run.font.color.rgb = properties['font']['color']
        run.style = properties['style']
        
        # Safely restore XML-level properties if they exist
        try:
            if hasattr(run._element, 'rPr'):
                if 'spacing' in properties and properties['spacing'] is not None:
                    run._element.rPr.spacing = properties['spacing']
                if 'position' in properties and properties['position'] is not None:
                    run._element.rPr.position = properties['position']
        except:
            pass

    async def process_paragraph(self, paragraph, session, target_language: str):
        """Process a paragraph while preserving exact formatting and inline images"""
        if not paragraph.text.strip():
            return

        # Store original properties for each run
        original_properties = []
        original_texts = []
        has_inline_images = False
        
        # Collect all runs and their properties, preserving inline shapes
        for run in paragraph.runs:
            # Check for inline shapes (images) in the run
            if hasattr(run, '_r') and run._r.drawing_lst:
                has_inline_images = True
                original_properties.append({
                    'properties': self.preserve_run_properties(run),
                    'has_image': True,
                    'image_data': run._r.drawing_lst
                })
            else:
                original_properties.append({
                    'properties': self.preserve_run_properties(run),
                    'has_image': False
                })
            original_texts.append(run.text)

        # If paragraph contains images, preserve the entire run structure
        if has_inline_images:
            try:
                # Only translate text runs, preserve image runs exactly
                for i, (props, original) in enumerate(zip(original_properties, original_texts)):
                    if not props['has_image'] and original.strip():
                        # Translate only text content
                        translated_text = await self.translate_chunk_with_retry(session, original, target_language)
                        translated_text = self.ensure_proper_spacing(translated_text)
                        # Preserve original spacing
                        if original.startswith(" "):
                            translated_text = " " + translated_text
                        if original.endswith(" "):
                            translated_text = translated_text + " "
                        paragraph.runs[i].text = translated_text
                    # Restore all original properties
                    self.restore_run_properties(paragraph.runs[i], props['properties'])
            except Exception as e:
                print(f"Error processing paragraph with images: {str(e)}")
                # Restore original text and properties on error
                for i, (props, text) in enumerate(zip(original_properties, original_texts)):
                    paragraph.runs[i].text = text
                    self.restore_run_properties(paragraph.runs[i], props['properties'])
            return

        # For paragraphs without images, proceed with normal translation
        try:
            # Translate the entire paragraph text to maintain context
            full_text = paragraph.text
            translated_text = await self.translate_chunk_with_retry(session, full_text, target_language)
            
            # Clean up translation while preserving intentional formatting
            translated_text = self.ensure_proper_spacing(translated_text)
            
            # Split translated text proportionally among runs
            words = translated_text.split()
            total_original_length = sum(len(text) for text in original_texts)
            
            current_word = 0
            for i, (props, original) in enumerate(zip(original_properties, original_texts)):
                if not words[current_word:]:
                    paragraph.runs[i].text = ""
                    continue
                    
                # Calculate how many words should go in this run
                run_proportion = len(original) / total_original_length if total_original_length > 0 else 0
                words_for_run = max(1, int(run_proportion * len(words)))
                
                # Get text for this run
                run_text = " ".join(words[current_word:current_word + words_for_run])
                current_word += words_for_run
                
                # Preserve original spacing
                if original.startswith(" "):
                    run_text = " " + run_text
                if original.endswith(" "):
                    run_text = run_text + " "
                
                # Update run text and restore properties
                paragraph.runs[i].text = run_text
                self.restore_run_properties(paragraph.runs[i], props['properties'])
            
            # Add any remaining words to the last run
            if current_word < len(words) and paragraph.runs:
                last_text = paragraph.runs[-1].text
                last_text += " " + " ".join(words[current_word:])
                paragraph.runs[-1].text = last_text
                
        except Exception as e:
            print(f"Error processing paragraph: {str(e)}")
            # Restore original text and properties on error
            for i, (props, text) in enumerate(zip(original_properties, original_texts)):
                paragraph.runs[i].text = text
                self.restore_run_properties(paragraph.runs[i], props['properties'])
